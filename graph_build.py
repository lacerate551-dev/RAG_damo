"""
Graph Build - 图谱构建工具（轻量级版本）

功能：
1. 指定单个或多个文件构建图谱
2. 增量添加到现有图谱（不删除已有数据）
3. 清空并重建图谱
4. 查看图谱统计信息

使用方式：
    python graph_build.py --help
    python graph_build.py --stats
    python graph_build.py --file documents/差旅管理办法.pdf
    python graph_build.py --files documents/制度1.pdf documents/制度2.docx
    python graph_build.py --dir documents/人事制度
    python graph_build.py --clear --file documents/差旅管理办法.pdf
    python graph_build.py -i  # 交互模式
"""

import os
import sys
import json
import re
import argparse
from dataclasses import dataclass
from typing import List, Dict, Any, Optional

# Windows 控制台编码处理
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# 添加项目路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 只导入必要的配置
try:
    from config import (
        API_KEY, BASE_URL,
        NEO4J_URI, NEO4J_USER, NEO4J_PASSWORD,
        USE_GRAPH_RAG, GRAPH_EXTRACTION_MODEL
    )
except ImportError as e:
    print(f"错误: 配置导入失败 - {e}")
    print("请确保 config.py 存在并包含必要配置")
    sys.exit(1)

from openai import OpenAI

# 导入图谱管理器
try:
    from graph_manager import GraphManager, Entity, Triple
except ImportError as e:
    print(f"错误: 无法导入 graph_manager - {e}")
    sys.exit(1)


# ==================== 文档解析（不依赖向量模型） ====================

def extract_text_from_txt(filepath: str) -> str:
    """从TXT文件提取文本"""
    encodings = ['utf-8', 'gbk', 'gb2312']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_from_pdf(filepath: str) -> List[Dict]:
    """从PDF提取文本"""
    pages = []
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append({
                        'text': text,
                        'page': page_num + 1
                    })
    except ImportError:
        print("警告: pdfplumber 未安装，无法解析 PDF")
    except Exception as e:
        print(f"PDF 解析错误: {e}")
    return pages


def extract_text_from_docx(filepath: str) -> List[Dict]:
    """从Word文档提取文本"""
    blocks = []
    try:
        from docx import Document
        doc = Document(filepath)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append({'text': text, 'is_table': False})
        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    blocks.append({'text': row_text, 'is_table': True})
    except ImportError:
        print("警告: python-docx 未安装，无法解析 Word 文档")
    except Exception as e:
        print(f"Word 解析错误: {e}")
    return blocks


def extract_text_from_xlsx(filepath: str) -> List[Dict]:
    """从Excel提取文本"""
    rows = []
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                text = " | ".join(str(cell) if cell else "" for cell in row)
                if text.strip():
                    rows.append({
                        'text': text,
                        'sheet': sheet_name
                    })
    except ImportError:
        print("警告: openpyxl 未安装，无法解析 Excel")
    except Exception as e:
        print(f"Excel 解析错误: {e}")
    return rows


def split_text(text: str, chunk_size: int = 2000) -> List[str]:
    """文本分块"""
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    chunks = []
    # 按段落分割
    paragraphs = text.split('\n\n')
    current_chunk = ""

    for p in paragraphs:
        if len(current_chunk) + len(p) > chunk_size:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = p
        else:
            current_chunk += "\n\n" + p

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


def extract_document(filepath: str) -> List[Dict]:
    """
    从文档提取文本块

    Returns:
        List of {'content': str, 'source': str, 'page': int, ...}
    """
    ext = os.path.splitext(filepath)[1].lower()
    filename = os.path.basename(filepath)
    chunks = []

    if ext == '.txt':
        text = extract_text_from_txt(filepath)
        for i, chunk in enumerate(split_text(text)):
            chunks.append({'content': chunk, 'source': filename, 'chunk': i})

    elif ext == '.pdf':
        pages = extract_text_from_pdf(filepath)
        for page in pages:
            for chunk in split_text(page['text']):
                chunks.append({
                    'content': chunk,
                    'source': filename,
                    'page': page['page']
                })

    elif ext in ['.docx', '.doc']:
        blocks = extract_text_from_docx(filepath)
        for i, block in enumerate(blocks):
            if len(block['text']) > 50:
                chunks.append({
                    'content': block['text'],
                    'source': filename,
                    'chunk': i,
                    'is_table': block.get('is_table', False)
                })

    elif ext == '.xlsx':
        rows = extract_text_from_xlsx(filepath)
        for i, row in enumerate(rows):
            if len(row['text']) > 20:
                chunks.append({
                    'content': row['text'],
                    'source': filename,
                    'sheet': row['sheet'],
                    'row': i
                })

    return chunks


# ==================== 实体提取器（只依赖 LLM） ====================

# 实体类型定义
ENTITY_TYPES = {
    "部门": "组织机构，如：人力资源部、财务部、行政部、信息技术部",
    "制度": "规章制度，如：差旅管理办法、报销制度、信息安全管理制度",
    "人员": "人员角色，如：员工、经理、审批人、部门负责人、总监",
    "流程": "业务流程，如：报销流程、审批流程、申请流程、入职流程",
    "金额": "金额标准，如：补助金额、报销限额、费用标准",
    "时间": "时间期限，如：有效期、申请时限、审批时限",
    "条件": "适用条件，如：享受条件、申请条件、适用范围",
    "地点": "地点场所，如：出差地点、办公地点",
    "信息": "信息类型，如：绝密级信息、机密级信息、秘密级信息",
    "项目": "业务项目，如：培训项目、工程项目",
}

# 关系类型定义
RELATION_TYPES = {
    "负责": "部门/人员 对 制度/流程 的管理责任",
    "适用": "制度 对 人员/部门 的适用范围",
    "包含": "制度/流程 包含的子项",
    "审批": "人员 对 流程 的审批权限",
    "限额": "制度 规定的 金额 限制",
    "时效": "制度 规定的 时间 限制",
    "条件": "制度 规定的 适用条件",
    "相关": "制度 与 其他制度 的关联",
    "属于": "人员/制度 属于 某个部门",
    "管理": "部门/人员 管理的范围",
}


class EntityExtractor:
    """实体提取器（只使用 LLM，不依赖向量模型）"""

    def __init__(self, model: str = None):
        self.model = model or GRAPH_EXTRACTION_MODEL
        self.client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
        self.entity_types = ENTITY_TYPES
        self.relation_types = RELATION_TYPES

    def extract(self, text: str, doc_source: str = None) -> List[Triple]:
        """
        从文本提取实体和关系

        Args:
            text: 输入文本
            doc_source: 文档来源

        Returns:
            三元组列表
        """
        # 限制文本长度
        if len(text) > 2500:
            text = text[:2500] + "..."

        entity_desc = "\n".join([f"- {k}: {v}" for k, v in self.entity_types.items()])
        relation_desc = "\n".join([f"- {k}: {v}" for k, v in self.relation_types.items()])

        prompt = f"""从以下企业制度文本中提取实体和关系。

## 实体类型
{entity_desc}

## 关系类型
{relation_desc}

## 待提取文本
{text}

## 输出要求
请输出JSON格式，包含entities和relations两个数组：
- entities: 每个实体包含 name, type, properties(可选)
- relations: 每个关系包含 head(头实体名), relation(关系类型), tail(尾实体名)

只输出JSON，不要其他内容："""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "你是知识图谱构建专家，擅长从制度文档中提取实体和关系。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )

            raw_response = response.choices[0].message.content
            return self._parse_response(raw_response, doc_source)

        except Exception as e:
            print(f"    提取失败: {e}")
            return []

    def _parse_response(self, response: str, doc_source: str = None) -> List[Triple]:
        """解析 LLM 响应"""
        triples = []
        try:
            # 清理响应
            json_str = response.strip()
            json_match = re.search(r'\{[\s\S]*\}', json_str)
            if json_match:
                json_str = json_match.group(0)

            data = json.loads(json_str)

            # 构建实体映射
            entity_map = {}
            for e in data.get('entities', []):
                name = e.get('name', '').strip()
                etype = e.get('type', '').strip()
                props = e.get('properties', {}) or {}
                if name and etype:
                    if doc_source:
                        props['source'] = doc_source
                    entity_map[name] = Entity(name=name, type=etype, properties=props)

            # 构建三元组
            for r in data.get('relations', []):
                head = r.get('head', '').strip()
                relation = r.get('relation', '').strip()
                tail = r.get('tail', '').strip()

                if head and relation and tail:
                    head_e = entity_map.get(head) or Entity(name=head, type="未知")
                    tail_e = entity_map.get(tail) or Entity(name=tail, type="未知")

                    if relation in self.relation_types:
                        triples.append(Triple(head=head_e, relation=relation, tail=tail_e))

        except Exception as e:
            print(f"    解析失败: {e}")

        return triples

    def extract_batch(self, texts: List[str], sources: List[str] = None, verbose: bool = True) -> List[Triple]:
        """批量提取"""
        all_triples = []
        sources = sources or [None] * len(texts)

        for i, (text, source) in enumerate(zip(texts, sources)):
            if verbose:
                print(f"  处理 [{i+1}/{len(texts)}]...", end=" ")

            triples = self.extract(text, source)
            all_triples.extend(triples)

            if verbose:
                print(f"提取 {len(triples)} 个三元组")

        # 去重
        seen = set()
        unique = []
        for t in all_triples:
            key = (t.head.name, t.relation, t.tail.name)
            if key not in seen:
                seen.add(key)
                unique.append(t)

        return unique


# ==================== 图谱构建函数 ====================

DOCUMENTS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "documents")


def build_graph_from_files(filepaths: List[str], clear_existing: bool = False, verbose: bool = True) -> int:
    """
    从文件列表构建图谱

    Args:
        filepaths: 文件路径列表
        clear_existing: 是否清空现有图谱
        verbose: 是否打印详细信息

    Returns:
        成功写入的三元组数量
    """
    # 连接 Neo4j
    gm = GraphManager()
    if not gm.connect():
        print("错误: 无法连接到 Neo4j")
        return 0

    # 清空现有图谱
    if clear_existing:
        if verbose:
            print("清空现有图谱...")
        gm.clear_graph()

    # 创建提取器
    extractor = EntityExtractor()

    all_triples = []

    for filepath in filepaths:
        if not os.path.exists(filepath):
            if verbose:
                print(f"警告: 文件不存在 {filepath}")
            continue

        filename = os.path.basename(filepath)
        if verbose:
            print(f"\n处理文件: {filename}")

        # 提取文档内容
        chunks = extract_document(filepath)
        if not chunks:
            if verbose:
                print(f"  无有效内容")
            continue

        if verbose:
            print(f"  文本块数: {len(chunks)}")

        # 提取实体和关系
        texts = [c['content'] for c in chunks]
        sources = [c['source'] for c in chunks]

        triples = extractor.extract_batch(texts, sources, verbose=verbose)
        all_triples.extend(triples)

    if not all_triples:
        if verbose:
            print("\n未提取到任何三元组")
        gm.close()
        return 0

    # 去重
    seen = set()
    unique_triples = []
    for t in all_triples:
        key = (t.head.name, t.relation, t.tail.name)
        if key not in seen:
            seen.add(key)
            unique_triples.append(t)

    if verbose:
        print(f"\n总计 {len(unique_triples)} 个唯一三元组")

    # 写入图谱
    if verbose:
        print("写入图谱...")

    success = gm.build_from_triples(unique_triples, verbose=verbose)

    # 获取统计
    stats = gm.get_stats()
    if verbose:
        print(f"\n图谱统计:")
        print(f"  节点数: {stats['nodes']}")
        print(f"  关系数: {stats['edges']}")
        if stats['types']:
            print(f"  实体类型:")
            for t, c in sorted(stats['types'].items(), key=lambda x: -x[1])[:10]:
                print(f"    - {t}: {c}")

    gm.close()
    return success


def build_graph_from_directory(directory: str, clear_existing: bool = False, verbose: bool = True) -> int:
    """从目录构建图谱"""
    if not os.path.isdir(directory):
        print(f"错误: 目录不存在 {directory}")
        return 0

    # 收集支持的文件
    supported_extensions = {'.pdf', '.docx', '.doc', '.xlsx', '.txt'}
    filepaths = []

    for root, dirs, files in os.walk(directory):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()
            if ext in supported_extensions:
                filepath = os.path.join(root, filename)
                filepaths.append(filepath)

    if not filepaths:
        print(f"警告: 目录中没有支持的文件 {directory}")
        return 0

    if verbose:
        print(f"\n找到 {len(filepaths)} 个文件:")
        for fp in filepaths[:10]:
            print(f"  - {os.path.basename(fp)}")
        if len(filepaths) > 10:
            print(f"  ... 还有 {len(filepaths) - 10} 个文件")

    return build_graph_from_files(filepaths, clear_existing, verbose)


def show_stats():
    """显示图谱统计"""
    gm = GraphManager()
    if not gm.connect():
        print("错误: 无法连接到 Neo4j")
        return

    stats = gm.get_stats()
    print("\n图谱统计:")
    print(f"  节点数: {stats['nodes']}")
    print(f"  关系数: {stats['edges']}")
    if stats['types']:
        print(f"  实体类型:")
        for t, c in sorted(stats['types'].items(), key=lambda x: -x[1]):
            print(f"    - {t}: {c}")

    gm.close()


def interactive_mode():
    """交互模式"""
    print("\n" + "=" * 60)
    print("Graph Build - 图谱构建工具")
    print("=" * 60)

    # 检查连接
    gm = GraphManager()
    if not gm.connect():
        print("错误: 无法连接到 Neo4j")
        return

    stats = gm.get_stats()
    print(f"\n当前图谱: {stats['nodes']} 节点, {stats['edges']} 关系")
    gm.close()

    print("\n命令:")
    print("  stats          - 查看统计")
    print("  add <文件>     - 添加文件")
    print("  dir <目录>     - 添加目录")
    print("  clear          - 清空图谱")
    print("  rebuild <文件> - 清空并重建")
    print("  quit           - 退出")
    print("=" * 60)

    while True:
        try:
            cmd = input("\n> ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n再见!")
            break

        if not cmd:
            continue

        parts = cmd.split(maxsplit=1)
        action = parts[0].lower()
        arg = parts[1] if len(parts) > 1 else None

        if action == 'quit':
            print("再见!")
            break
        elif action == 'stats':
            show_stats()
        elif action == 'clear':
            confirm = input("确认清空图谱? (yes/no): ").strip().lower()
            if confirm == 'yes':
                gm = GraphManager()
                if gm.connect():
                    gm.clear_graph()
                    gm.close()
                    print("图谱已清空")
        elif action == 'add' and arg:
            filepath = arg if os.path.isabs(arg) else os.path.join(DOCUMENTS_PATH, arg)
            if os.path.exists(filepath):
                build_graph_from_files([filepath], clear_existing=False)
            else:
                print(f"文件不存在: {arg}")
        elif action == 'dir' and arg:
            directory = arg if os.path.isabs(arg) else os.path.join(DOCUMENTS_PATH, arg)
            if os.path.isdir(directory):
                build_graph_from_directory(directory, clear_existing=False)
            else:
                print(f"目录不存在: {arg}")
        elif action == 'rebuild' and arg:
            filepath = arg if os.path.isabs(arg) else os.path.join(DOCUMENTS_PATH, arg)
            if os.path.exists(filepath):
                build_graph_from_files([filepath], clear_existing=True)
            else:
                print(f"文件不存在: {arg}")
        else:
            print(f"未知命令: {action}")
            print("输入 'quit' 退出")


def main():
    parser = argparse.ArgumentParser(
        description='Graph Build - 图谱构建工具（轻量级版本）',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s --stats                    # 查看统计
  %(prog)s --file documents/文件.txt   # 从文件构建
  %(prog)s --dir documents/制度        # 从目录构建
  %(prog)s --clear --file 文件.txt     # 清空并重建
  %(prog)s -i                         # 交互模式
        """
    )

    parser.add_argument('--stats', action='store_true', help='查看图谱统计')
    parser.add_argument('--file', type=str, help='从单个文件构建')
    parser.add_argument('--files', nargs='+', help='从多个文件构建')
    parser.add_argument('--dir', type=str, help='从目录构建')
    parser.add_argument('--clear', action='store_true', help='构建前清空图谱')
    parser.add_argument('-i', '--interactive', action='store_true', help='交互模式')

    args = parser.parse_args()

    # 检查 Graph RAG 是否启用
    if not USE_GRAPH_RAG:
        print("警告: Graph RAG 功能未启用 (USE_GRAPH_RAG=False)")

    # 查看统计
    if args.stats:
        show_stats()
        return

    # 交互模式
    if args.interactive:
        interactive_mode()
        return

    # 从单个文件构建
    if args.file:
        filepath = args.file if os.path.isabs(args.file) else os.path.join(DOCUMENTS_PATH, args.file)
        if not os.path.exists(filepath):
            print(f"错误: 文件不存在 {filepath}")
            return
        build_graph_from_files([filepath], clear_existing=args.clear)
        return

    # 从多个文件构建
    if args.files:
        filepaths = []
        for fp in args.files:
            full_path = fp if os.path.isabs(fp) else os.path.join(DOCUMENTS_PATH, fp)
            if os.path.exists(full_path):
                filepaths.append(full_path)
            else:
                print(f"警告: 文件不存在 {fp}")
        if filepaths:
            build_graph_from_files(filepaths, clear_existing=args.clear)
        return

    # 从目录构建
    if args.dir:
        directory = args.dir if os.path.isabs(args.dir) else os.path.join(DOCUMENTS_PATH, args.dir)
        if not os.path.isdir(directory):
            print(f"错误: 目录不存在 {directory}")
            return
        build_graph_from_directory(directory, clear_existing=args.clear)
        return

    # 无参数显示帮助
    parser.print_help()


if __name__ == "__main__":
    main()

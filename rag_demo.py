"""
RAG Demo - 基于本地向量模型 + Chroma + Qwen API 的简单知识库问答系统
支持格式: PDF, Word(.docx/.doc), Excel(.xlsx), TXT
混合检索: 向量检索 + BM25关键词检索 + Rerank重排序

模型目录结构:
  models/
  ├── bge-base-zh-v1.5/      # 向量模型（必需）
  │   ├── config.json
  │   ├── pytorch_model.bin
  │   ├── tokenizer.json
  │   └── vocab.txt
  └── bge-reranker-base/     # 重排序模型（首次运行自动下载）
      ├── config.json
      ├── model.safetensors
      └── tokenizer.json
"""

import os
import json
import pickle
from sentence_transformers import SentenceTransformer, CrossEncoder
import chromadb
from openai import OpenAI
import pdfplumber
from docx import Document
from openpyxl import load_workbook
import docx2txt
import numpy as np
from rank_bm25 import BM25Okapi
import jieba

# OpenDataLoader PDF 解析器（可选）
try:
    from parsers.pdf_odl import parse_pdf_with_odl, ChunkMetadata
    ODL_AVAILABLE = True
except ImportError:
    ODL_AVAILABLE = False

# 语义分块器（可选）
try:
    from core.chunker import SemanticChunker, HybridChunker
    SEMANTIC_CHUNKER_AVAILABLE = True
except ImportError:
    SEMANTIC_CHUNKER_AVAILABLE = False

# Docling 文档解析器（可选）
try:
    from parsers.docx_docling import DoclingParser, DOCLING_AVAILABLE, DocChunk
    if not DOCLING_AVAILABLE:
        DOCLING_AVAILABLE = False
except ImportError:
    DOCLING_AVAILABLE = False

# Excel 增强解析器（可选）
try:
    from parsers.excel_parser import ExcelParserEnhanced, ExcelChunk
    EXCEL_ENHANCED_AVAILABLE = True
except ImportError:
    EXCEL_ENHANCED_AVAILABLE = False

# 导入配置
try:
    from config import API_KEY, BASE_URL, MODEL
except ImportError:
    print("错误: 未找到config.py文件")
    print("请复制config.example.py为config.py并填入你的API Key")
    exit(1)

# ========== 路径配置 ==========
# 项目根目录
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))

# 模型目录（统一管理）
MODELS_DIR = os.path.join(PROJECT_ROOT, "models")

# 向量模型路径（必需，需手动下载）
EMBEDDING_MODEL_PATH = os.path.join(MODELS_DIR, "bge-base-zh-v1.5")

# 重排序模型路径（首次运行自动下载到此目录）
RERANK_MODEL_PATH = os.path.join(MODELS_DIR, "bge-reranker-base")

# 数据目录
VECTOR_STORE_PATH = os.path.join(PROJECT_ROOT, "vector_store")
CHROMA_DB_PATH = os.path.join(VECTOR_STORE_PATH, "chroma")
DOCUMENTS_PATH = os.path.join(PROJECT_ROOT, "documents")
BM25_INDEXES_PATH = os.path.join(VECTOR_STORE_PATH, "bm25")

# 多向量库配置
USE_MULTI_KB = True  # 是否使用多向量库模式

# 混合检索配置
USE_HYBRID_SEARCH = True  # 是否启用混合检索
VECTOR_WEIGHT = 0.5  # 向量检索权重
BM25_WEIGHT = 0.5  # BM25检索权重

# Rerank配置
USE_RERANK = True  # 是否启用重排序
RERANK_CANDIDATES = 20  # 重排序前的候选数量（混合检索后）
RERANK_TOP_K = 5  # 重排序后返回的数量

# PDF 解析配置
USE_ODL_PARSER = True  # 是否使用 OpenDataLoader PDF 解析器（更高质量）
ODL_USE_STRUCT_TREE = True  # 是否使用 PDF 结构树（Tagged PDF）
ODL_USE_HYBRID = False  # 是否使用混合模式（需要后端服务）

# Docling 文档解析配置
USE_DOCLING_PARSER = True  # 是否使用 Docling 解析 Word（更高质量）

# Excel 解析配置
USE_EXCEL_ENHANCED = True  # 是否使用增强版 Excel 解析器
EXCEL_MAX_ROWS_PER_CHUNK = 50  # 每个分块的最大行数
EXCEL_MIN_ROWS_PER_CHUNK = 2   # 每个分块的最小行数

# 分块配置
USE_SEMANTIC_CHUNK = True  # 是否使用语义分块（替代固定字符切分）
SEMANTIC_BREAKPOINT_THRESHOLD = 0.5  # 语义分块阈值（值越大分块越少）
SEMANTIC_MIN_CHUNK_SIZE = 50  # 最小分块字符数
SEMANTIC_MAX_CHUNK_SIZE = 800  # 最大分块字符数


# ========== 模型初始化辅助函数 ==========
def ensure_models_dir():
    """确保模型目录存在"""
    os.makedirs(MODELS_DIR, exist_ok=True)


def check_embedding_model():
    """检查向量模型是否存在"""
    required_files = ["config.json", "pytorch_model.bin", "tokenizer.json", "vocab.txt"]
    model_dir = EMBEDDING_MODEL_PATH

    if not os.path.exists(model_dir):
        return False, f"模型目录不存在: {model_dir}"

    missing_files = []
    for f in required_files:
        if not os.path.exists(os.path.join(model_dir, f)):
            missing_files.append(f)

    if missing_files:
        return False, f"缺少文件: {', '.join(missing_files)}"

    return True, None


# ========== BM25索引管理器 ==========
class BM25Index:
    """BM25索引管理器，用于关键词检索"""

    def __init__(self):
        self.bm25 = None
        self.documents = []  # 原始文档
        self.metadatas = []  # 元数据
        self.ids = []  # 文档ID

    def tokenize(self, text):
        """中文分词"""
        return list(jieba.cut(text))

    def add_documents(self, ids, documents, metadatas):
        """添加文档到索引"""
        self.ids = ids
        self.documents = documents
        self.metadatas = metadatas

        # 分词并构建BM25索引
        tokenized_docs = [self.tokenize(doc) for doc in documents]
        self.bm25 = BM25Okapi(tokenized_docs)

    def search(self, query, top_k=10):
        """BM25检索"""
        if not self.bm25:
            return {'ids': [[]], 'documents': [[]], 'metadatas': [[]], 'distances': [[]]}

        tokenized_query = self.tokenize(query)
        scores = self.bm25.get_scores(tokenized_query)

        # 获取top_k个结果
        top_indices = np.argsort(scores)[::-1][:top_k]

        return {
            'ids': [[self.ids[i] for i in top_indices]],
            'documents': [[self.documents[i] for i in top_indices]],
            'metadatas': [[self.metadatas[i] for i in top_indices]],
            'distances': [[float(scores[i]) for i in top_indices]]
        }

    def save(self, path):
        """保存索引到文件"""
        data = {
            'ids': self.ids,
            'documents': self.documents,
            'metadatas': self.metadatas
        }
        with open(path, 'wb') as f:
            pickle.dump(data, f)
        print(f"      BM25索引已保存: {path}")

    def load(self, path):
        """从文件加载索引"""
        if not os.path.exists(path):
            return False

        with open(path, 'rb') as f:
            data = pickle.load(f)

        self.ids = data['ids']
        self.documents = data['documents']
        self.metadatas = data['metadatas']

        # 重建BM25索引
        tokenized_docs = [self.tokenize(doc) for doc in self.documents]
        self.bm25 = BM25Okapi(tokenized_docs)

        print(f"      BM25索引已加载: {len(self.documents)} 个文档")
        return True

    def clear(self):
        """清空索引"""
        self.bm25 = None
        self.documents = []
        self.metadatas = []
        self.ids = []


def rebuild_bm25_index():
    """从ChromaDB重建BM25索引"""
    # 获取所有文档
    results = collection.get()
    if not results['ids']:
        bm25_index.clear()
        return 0

    # 重建索引
    bm25_index.add_documents(
        ids=results['ids'],
        documents=results['documents'],
        metadatas=results['metadatas']
    )
    return len(results['ids'])


# ========== 初始化组件 ==========
print("=" * 50)
print("RAG Demo 知识库问答系统 (混合检索版)")
print("=" * 50)

# 确保模型目录存在
ensure_models_dir()

# [1/5] 加载本地向量模型
print("\n[1/5] 加载本地向量模型...")
model_ok, model_error = check_embedding_model()
if not model_ok:
    print(f"\n错误: 向量模型未正确安装!")
    print(f"  {model_error}")
    print("\n请按以下步骤下载模型:")
    print("  1. 创建模型目录: mkdir models")
    print("  2. 下载向量模型:")
    print("     pip install huggingface-hub")
    print("     huggingface-cli download BAAI/bge-base-zh-v1.5 --local-dir ./models/bge-base-zh-v1.5")
    print("\n  或者手动下载: https://huggingface.co/BAAI/bge-base-zh-v1.5")
    print("  将文件放入: ./models/bge-base-zh-v1.5/")
    exit(1)

embedding_model = SentenceTransformer(EMBEDDING_MODEL_PATH)
print(f"      模型加载完成: {EMBEDDING_MODEL_PATH}")

# 初始化语义分块器
semantic_chunker = None
if USE_SEMANTIC_CHUNK and SEMANTIC_CHUNKER_AVAILABLE:
    semantic_chunker = SemanticChunker(
        embedding_model=embedding_model,
        breakpoint_threshold=SEMANTIC_BREAKPOINT_THRESHOLD,
        min_chunk_size=SEMANTIC_MIN_CHUNK_SIZE,
        max_chunk_size=SEMANTIC_MAX_CHUNK_SIZE
    )
    print(f"      语义分块器已启用 (阈值: {SEMANTIC_BREAKPOINT_THRESHOLD})")
elif USE_SEMANTIC_CHUNK and not SEMANTIC_CHUNKER_AVAILABLE:
    print("      警告: 语义分块模块未找到，使用传统分块")

print("\n[2/5] 初始化向量数据库...")
# 多向量库模式
if USE_MULTI_KB:
    from knowledge.manager import KnowledgeBaseManager
    from knowledge.router import KnowledgeBaseRouter
    from auth.gateway import get_accessible_collections

    kb_manager = KnowledgeBaseManager(CHROMA_DB_PATH)
    kb_router = KnowledgeBaseRouter(use_llm=False)
    print(f"      数据库路径: {CHROMA_DB_PATH}")
    print(f"      向量库数量: {len(kb_manager.list_collections())}")
    print(f"      多向量库模式: 已启用")

    # 兼容旧代码的 collection 引用（指向 public_kb）
    chroma_client = None
    collection = kb_manager.get_collection('public_kb')
else:
    chroma_client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
    collection = chroma_client.get_or_create_collection(
        name="knowledge_base",
        metadata={"description": "RAG Demo 知识库"}
    )
    kb_manager = None
    kb_router = None
    print(f"      数据库路径: {CHROMA_DB_PATH}")
print(f"      当前模式: {'多向量库' if USE_MULTI_KB else '单向量库'}")

print("\n[3/5] 初始化BM25索引...")
bm25_index = BM25Index()
if USE_HYBRID_SEARCH:
    if USE_MULTI_KB:
        print("      多向量库模式: BM25索引按需从各向量库加载")
    else:
        bm25_index.load(BM25_INDEX_PATH)
else:
    print("      混合检索已禁用，跳过BM25索引加载")

print("\n[4/5] 初始化大模型客户端...")
llm_client = OpenAI(
    api_key=API_KEY,
    base_url=BASE_URL
)
print(f"      API地址: {BASE_URL}")
print(f"      模型: {MODEL}")

# 初始化Reranker模型
reranker = None
if USE_RERANK:
    print("\n[5/5] 加载重排序模型...")
    try:
        # 检查本地模型是否存在
        if os.path.exists(RERANK_MODEL_PATH) and os.path.exists(os.path.join(RERANK_MODEL_PATH, "config.json")):
            # 使用本地模型
            reranker = CrossEncoder(RERANK_MODEL_PATH)
            print(f"      Rerank模型加载完成: {RERANK_MODEL_PATH}")
        else:
            # 从 HuggingFace 下载并指定缓存目录
            print(f"      首次运行，正在下载 Rerank 模型...")
            print(f"      模型将保存到: {RERANK_MODEL_PATH}")
            os.makedirs(RERANK_MODEL_PATH, exist_ok=True)

            # 下载模型
            from transformers import AutoModelForSequenceClassification, AutoTokenizer
            model = AutoModelForSequenceClassification.from_pretrained("BAAI/bge-reranker-base")
            tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-reranker-base")

            # 保存到本地
            model.save_pretrained(RERANK_MODEL_PATH)
            tokenizer.save_pretrained(RERANK_MODEL_PATH)

            # 加载模型
            reranker = CrossEncoder(RERANK_MODEL_PATH)
            print(f"      Rerank模型下载完成: {RERANK_MODEL_PATH}")
    except Exception as e:
        print(f"      Rerank模型加载失败: {e}")
        print("      将使用纯向量检索模式（不影响基本功能）")
        print("      如需使用 Rerank，请手动下载模型到 ./models/bge-reranker-base/")
        USE_RERANK = False
else:
    print("\n[5/5] 跳过重排序模型加载 (已禁用)")


# ========== 文件管理函数 ==========
def list_indexed_files():
    """列出向量库中已索引的文件"""
    results = collection.get()
    if not results['metadatas']:
        return {}

    # 统计每个文件的片段数
    from collections import Counter
    file_chunks = Counter()
    for meta in results['metadatas']:
        file_chunks[meta['source']] += 1

    return dict(file_chunks)


def delete_file_from_index(filename):
    """从向量库中删除指定文件的所有片段"""
    # 获取该文件的所有ID
    results = collection.get(
        where={"source": filename}
    )

    if not results['ids']:
        print(f"文件未在知识库中: {filename}")
        return 0

    # 删除
    collection.delete(ids=results['ids'])
    deleted_count = len(results['ids'])
    print(f"已删除 {filename} 的 {deleted_count} 个片段")
    return deleted_count


def add_file_to_index(filepath):
    """添加单个文件到向量库"""
    rel_path = os.path.relpath(filepath, DOCUMENTS_PATH)
    ext = os.path.splitext(filepath)[1].lower()

    supported_extensions = {'.txt', '.pdf', '.docx', '.xlsx'}
    if ext not in supported_extensions:
        print(f"不支持的文件格式: {ext}")
        return 0

    total_chunks = 0

    # 根据文件类型处理
    if ext == '.pdf':
        pages = extract_text_from_pdf(filepath)
        if pages:
            for page_info in pages:
                page_text = page_info['text']
                page_num = page_info['page']
                has_table = page_info.get('has_table', False)
                section = page_info.get('section', '')

                # ODL 分块不再切分，传统 PDF 使用智能分块
                is_odl_chunk = page_info.get('is_odl_chunk', False)
                if is_odl_chunk:
                    chunks = [page_text]
                else:
                    chunks = smart_split_text(page_text)

                for i, chunk in enumerate(chunks):
                    vector = embedding_model.encode(chunk).tolist()
                    chunk_id = f"{rel_path}_p{page_num}_{i}"

                    collection.add(
                        ids=[chunk_id],
                        embeddings=[vector],
                        documents=[chunk],
                        metadatas=[{
                            'source': rel_path,
                            'page': page_num,
                            'chunk_index': i,
                            'has_table': has_table,
                            'section': section
                        }]
                    )
                    total_chunks += 1
            print(f"添加 {rel_path}: {total_chunks} 个片段 (PDF, {len(pages)}页)")

    elif ext == '.docx':
        blocks = extract_text_from_docx(filepath)
        if blocks:
            for block in blocks:
                text = block['text']
                if len(text.strip()) < 10:
                    continue

                is_table = block.get('is_table', False)
                section = block.get('section', '')

                # 表格不分块，普通文本使用智能分块
                chunks = [text] if is_table else smart_split_text(text)

                for i, chunk in enumerate(chunks):
                    vector = embedding_model.encode(chunk).tolist()
                    chunk_id = f"{rel_path}_{total_chunks}_{i}"

                    collection.add(
                        ids=[chunk_id],
                        embeddings=[vector],
                        documents=[chunk],
                        metadatas=[{
                            'source': rel_path,
                            'chunk_index': total_chunks,
                            'is_table': is_table,
                            'section': section
                        }]
                    )
                    total_chunks += 1
            tables_count = sum(1 for b in blocks if b.get('is_table'))
            print(f"添加 {rel_path}: {total_chunks} 个片段 (Word, {len(blocks)}段落)")

    elif ext == '.xlsx':
        rows = extract_text_from_xlsx(filepath)
        if rows:
            for row_info in rows:
                text = row_info['text']
                if len(text.strip()) < 5:
                    continue

                sheet = row_info['sheet']
                row_num = row_info['row']
                is_header = row_info.get('is_header', False)
                header = row_info.get('header', '')

                full_text = text
                if header and not is_header:
                    full_text = f"【表头: {header}】\n{text}"

                vector = embedding_model.encode(full_text).tolist()
                chunk_id = f"{rel_path}_{sheet}_{row_num}"

                collection.add(
                    ids=[chunk_id],
                    embeddings=[vector],
                    documents=[full_text],
                    metadatas=[{
                        'source': rel_path,
                        'sheet': sheet,
                        'row': row_num,
                        'is_header': is_header
                    }]
                )
                total_chunks += 1
            sheets = set(r['sheet'] for r in rows)
            print(f"添加 {rel_path}: {total_chunks} 个片段 (Excel, {len(sheets)}工作表)")

    elif ext == '.txt':
        content = extract_text_from_txt(filepath)
        if content.strip():
            chunks = smart_split_text(content)  # 使用智能分块
            for i, chunk in enumerate(chunks):
                vector = embedding_model.encode(chunk).tolist()
                chunk_id = f"{rel_path}_{i}"

                collection.add(
                    ids=[chunk_id],
                    embeddings=[vector],
                    documents=[chunk],
                    metadatas=[{
                        'source': rel_path,
                        'chunk_index': i
                    }]
                )
                total_chunks += 1
            print(f"添加 {rel_path}: {total_chunks} 个片段 (TXT)")

    return total_chunks


def sync_documents():
    """同步文档目录与向量库（增量更新）"""
    print("\n" + "=" * 50)
    print("同步文档")
    print("=" * 50)

    # 获取向量库中已有的文件
    indexed_files = set(list_indexed_files().keys())
    print(f"\n向量库中已有 {len(indexed_files)} 个文件")

    # 扫描文档目录
    current_files = set()
    supported_extensions = {'.txt', '.pdf', '.docx', '.xlsx'}

    for root, dirs, files in os.walk(DOCUMENTS_PATH):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()
            if ext in supported_extensions:
                filepath = os.path.join(root, filename)
                rel_path = os.path.relpath(filepath, DOCUMENTS_PATH)
                current_files.add(rel_path)

    print(f"文档目录中有 {len(current_files)} 个文件")

    # 需要新增的文件
    files_to_add = current_files - indexed_files
    # 需要删除的文件
    files_to_delete = indexed_files - current_files

    if not files_to_add and not files_to_delete:
        print("\n文档已是最新，无需同步")
        return

    # 删除不存在的文件
    if files_to_delete:
        print(f"\n需要删除 {len(files_to_delete)} 个文件:")
        for f in sorted(files_to_delete):
            print(f"  - {f}")
            delete_file_from_index(f)

    # 添加新文件
    if files_to_add:
        print(f"\n需要添加 {len(files_to_add)} 个新文件:")
        for f in sorted(files_to_add):
            print(f"  + {f}")
            filepath = os.path.join(DOCUMENTS_PATH, f)
            add_file_to_index(filepath)

    print(f"\n向量库同步完成，当前共 {collection.count()} 个片段")

    # 重建BM25索引（增量更新后需要重建以保持一致）
    if USE_HYBRID_SEARCH:
        print("\n重建BM25索引...")
        rebuild_bm25_index()
        bm25_index.save(BM25_INDEX_PATH)
        print(f"BM25索引更新完成: {len(bm25_index.documents)} 个文档")


# ========== 文档处理函数 ==========
def extract_text_from_pdf(filepath):
    """
    从PDF提取文本，返回带页码和结构信息的内容列表

    支持两种解析模式：
    1. OpenDataLoader 模式（USE_ODL_PARSER=True）：高质量解析，保留文档结构
    2. pdfplumber 模式（默认）：基础解析，按页提取
    """
    # 使用 OpenDataLoader 解析器（优先）
    if USE_ODL_PARSER and ODL_AVAILABLE:
        return extract_text_from_pdf_odl(filepath)

    # 回退到 pdfplumber 解析
    return extract_text_from_pdf_plumber(filepath)


def extract_text_from_pdf_odl(filepath):
    """
    使用 OpenDataLoader PDF 解析器提取文本

    返回格式与 pdfplumber 兼容，但包含更丰富的结构信息：
    - 按章节分块而非按页
    - 保留标题层级
    - 包含 bounding box 坐标
    """
    pages_content = []
    try:
        result = parse_pdf_with_odl(
            filepath,
            use_struct_tree=ODL_USE_STRUCT_TREE,
            use_hybrid=ODL_USE_HYBRID
        )

        # 将分块结果转换为兼容格式
        for chunk in result['chunks']:
            pages_content.append({
                'text': chunk.content,
                'page': chunk.page_start,
                'page_end': chunk.page_end,
                'has_table': chunk.chunk_type == 'table',
                'section': chunk.title,
                'section_path': chunk.section_path,
                'level': chunk.level,
                'bbox': chunk.bbox,
                'source_file': chunk.source_file,
                'is_odl_chunk': True  # 标记为 ODL 分块
            })

        # 如果没有分块，回退到 pdfplumber
        if not pages_content:
            print(f"      OpenDataLoader 未提取到内容，回退到 pdfplumber: {filepath}")
            return extract_text_from_pdf_plumber(filepath)

    except Exception as e:
        print(f"      OpenDataLoader 解析错误，回退到 pdfplumber: {e}")
        return extract_text_from_pdf_plumber(filepath)

    return pages_content


def extract_text_from_pdf_plumber(filepath):
    """从PDF提取文本，返回带页码和结构信息的内容列表（pdfplumber 实现）"""
    pages_content = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    # 检测是否包含表格
                    tables = page.extract_tables()
                    has_table = len(tables) > 0

                    # 尝试识别章节标题（通常是短行、字号较大、以数字开头等）
                    lines = page_text.split('\n')
                    section_title = ""
                    for line in lines[:5]:  # 只检查前几行
                        line = line.strip()
                        # 章节标题特征：短、以数字或第X章/节开头
                        if len(line) < 30 and (line.startswith(('第', '一、', '二、', '三、', '四、', '五、', '1.', '2.', '3.')) or
                            any(keyword in line for keyword in ['章', '节', '规定', '制度', '办法'])):
                            section_title = line
                            break

                    pages_content.append({
                        'text': page_text,
                        'page': page_num + 1,
                        'has_table': has_table,
                        'section': section_title
                    })
    except Exception as e:
        print(f"      PDF解析错误 {filepath}: {e}")
    return pages_content


def extract_text_from_docx(filepath):
    """
    从Word文档提取文本，返回带结构信息的内容块列表

    支持两种解析模式：
    1. Docling 模式（USE_DOCLING_PARSER=True）：高质量解析，保留结构
    2. docx2txt + python-docx 模式（默认）：基础解析
    """
    # 使用 Docling 解析器（优先）
    if USE_DOCLING_PARSER and DOCLING_AVAILABLE:
        return extract_text_from_docx_docling(filepath)

    # 回退到传统解析
    return extract_text_from_docx_traditional(filepath)


def extract_text_from_docx_docling(filepath):
    """使用 Docling 解析 Word 文档"""
    content_blocks = []
    try:
        from parsers.docx_docling import DoclingParser

        parser = DoclingParser()
        result = parser.parse(filepath)

        # 转换分块为兼容格式
        for chunk in result['chunks']:
            content_blocks.append({
                'text': chunk.content,
                'is_heading': chunk.level > 0,
                'section': chunk.title,
                'section_path': chunk.section_path,
                'level': chunk.level,
                'is_table': chunk.chunk_type == 'table',
                'is_docling_chunk': True  # 标记为 Docling 分块
            })

        if not content_blocks:
            print(f"      Docling 未提取到内容，回退到传统解析: {filepath}")
            return extract_text_from_docx_traditional(filepath)

    except Exception as e:
        print(f"      Docling 解析错误，回退到传统解析: {e}")
        return extract_text_from_docx_traditional(filepath)

    return content_blocks


def extract_text_from_docx_traditional(filepath):
    """
    Word 传统解析（简化版，作为 Docling 的回退）

    保留基础文本提取和简单表格处理
    """
    content_blocks = []

    # 首先尝试用 docx2txt 提取（兼容性更好）
    try:
        full_text = docx2txt.process(filepath)
        if full_text and full_text.strip():
            for para in full_text.split('\n\n'):
                text = para.strip()
                if text:
                    content_blocks.append({
                        'text': text,
                        'is_heading': False,
                        'section': text[:30] if len(text) < 30 else text[:30] + '...',
                        'is_table': False
                    })
            return content_blocks
    except Exception as e:
        print(f"      docx2txt 解析失败，尝试 python-docx: {e}")

    # 备用：使用 python-docx
    try:
        doc = Document(filepath)
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 简单标题检测
            is_heading = para.style.name.startswith(('Heading', '标题'))
            if is_heading:
                current_section = text

            content_blocks.append({
                'text': text,
                'is_heading': is_heading,
                'section': current_section if current_section else text[:20],
                'is_table': False
            })

        # 简单表格提取
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    table_rows.append(row_text)
            if table_rows:
                content_blocks.append({
                    'text': "【表格】\n" + "\n".join(table_rows),
                    'is_heading': False,
                    'section': current_section,
                    'is_table': True
                })

    except Exception as e:
        print(f"      Word 解析错误 {filepath}: {e}")
    return content_blocks


def extract_text_from_xlsx(filepath):
    """
    从Excel提取文本，智能分块存储

    支持两种解析模式：
    1. 增强模式（USE_EXCEL_ENHANCED=True）：表头检测、合并单元格、智能分块
    2. 传统模式（默认）：简单的第一列/第二列规则分块
    """
    # 使用增强解析器（优先）
    if USE_EXCEL_ENHANCED and EXCEL_ENHANCED_AVAILABLE:
        return extract_text_from_xlsx_enhanced(filepath)

    # 回退到传统解析
    return extract_text_from_xlsx_traditional(filepath)


def extract_text_from_xlsx_enhanced(filepath):
    """使用增强解析器处理 Excel"""
    content_blocks = []
    try:
        from parsers.excel_parser import ExcelParserEnhanced

        parser = ExcelParserEnhanced(
            max_rows_per_chunk=EXCEL_MAX_ROWS_PER_CHUNK,
            min_rows_per_chunk=EXCEL_MIN_ROWS_PER_CHUNK
        )
        result = parser.parse(filepath)

        # 转换分块为兼容格式
        for chunk in result['chunks']:
            content_blocks.append({
                'text': chunk.content,
                'sheet': chunk.sheet,
                'row': int(chunk.row_range.split('-')[0]) if chunk.row_range else 0,
                'row_range': chunk.row_range,
                'col_range': chunk.col_range,
                'is_header': chunk.chunk_type == 'header',
                'block_title': chunk.title,
                'is_block': chunk.chunk_type == 'data',
                'headers': chunk.headers,
                'is_enhanced_chunk': True  # 标记为增强分块
            })

        if not content_blocks:
            print(f"      增强解析未提取到内容，回退到传统解析: {filepath}")
            return extract_text_from_xlsx_traditional(filepath)

    except Exception as e:
        print(f"      增强解析错误，回退到传统解析: {e}")
        return extract_text_from_xlsx_traditional(filepath)

    return content_blocks


def extract_text_from_xlsx_traditional(filepath):
    """
    Excel 传统解析（简化版，作为增强解析的回退）

    仅保留基础行读取，不做复杂的块检测
    """
    content_blocks = []
    try:
        wb = load_workbook(filepath, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = [str(cell) if cell is not None else "" for cell in row]
                row_text = " | ".join(cells)

                if row_text.strip(" |"):
                    content_blocks.append({
                        'text': row_text,
                        'sheet': sheet_name,
                        'row': row_idx,
                        'row_range': str(row_idx),
                        'is_header': row_idx == 1,
                        'block_title': '',
                        'is_block': False
                    })

    except Exception as e:
        print(f"      Excel解析错误 {filepath}: {e}")
    return content_blocks


def extract_text_from_txt(filepath):
    """从TXT提取文本"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            with open(filepath, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            print(f"      TXT解析错误 {filepath}: {e}")
            return ""
    except Exception as e:
        print(f"      TXT解析错误 {filepath}: {e}")
        return ""


def load_documents():
    """递归加载文档目录下的所有支持的文件，根据目录名自动分配安全级别"""
    documents = []
    supported_extensions = {'.txt', '.pdf', '.docx', '.doc', '.xlsx'}

    # 目录名 -> 安全级别映射 (精简版)
    security_levels = {
        'public': 'public',              # 全公司公开
        'confidential': 'confidential'   # 部门内保密
    }

    if not os.path.exists(DOCUMENTS_PATH):
        print(f"错误: 文档目录不存在 - {DOCUMENTS_PATH}")
        return documents

    # 递归遍历目录
    for root, dirs, files in os.walk(DOCUMENTS_PATH):
        for filename in files:
            filepath = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            if ext not in supported_extensions:
                continue

            # 计算相对路径（用于显示）
            rel_path = os.path.relpath(filepath, DOCUMENTS_PATH)

            # 根据目录路径确定安全级别
            security_level = 'internal'  # 默认部门内公开
            rel_path_lower = rel_path.lower().replace('\\', '/')
            for dir_name, level in security_levels.items():
                if f'/{dir_name}/' in f'/{rel_path_lower}/':
                    security_level = level
                    break

            # 根据文件类型提取文本
            if ext == '.pdf':
                # PDF特殊处理，返回带页码的列表
                pages = extract_text_from_pdf(filepath)
                if pages:
                    documents.append({
                        'filename': rel_path,
                        'type': 'pdf',
                        'pages': pages,
                        'security_level': security_level
                    })
                    print(f"      加载文档: {rel_path} (PDF, {len(pages)}页) [{security_level}]")
            elif ext in {'.docx', '.doc'}:
                # Word返回带结构的内容块列表
                blocks = extract_text_from_docx(filepath)
                if blocks:
                    documents.append({
                        'filename': rel_path,
                        'type': 'docx',
                        'blocks': blocks,
                        'security_level': security_level
                    })
                    tables_count = sum(1 for b in blocks if b.get('is_table'))
                    print(f"      加载文档: {rel_path} (Word, {len(blocks)}段落, {tables_count}表格) [{security_level}]")
            elif ext == '.xlsx':
                # Excel返回带行列信息的内容块列表
                blocks = extract_text_from_xlsx(filepath)
                if blocks:
                    documents.append({
                        'filename': rel_path,
                        'type': 'xlsx',
                        'rows': blocks,  # 兼容原结构
                        'security_level': security_level
                    })
                    multi_block_count = sum(1 for b in blocks if b.get('is_block'))
                    print(f"      加载文档: {rel_path} (Excel, {len(blocks)}个数据块, {multi_block_count}个多行块) [{security_level}]")
            elif ext == '.txt':
                content = extract_text_from_txt(filepath)
                if content.strip():
                    documents.append({
                        'filename': rel_path,
                        'content': content,
                        'type': 'txt',
                        'security_level': security_level
                    })
                    print(f"      加载文档: {rel_path} (TXT) [{security_level}]")

    return documents


def smart_split_text(text, force_semantic=False):
    """
    智能分块函数 - 根据配置选择分块策略

    Args:
        text: 待分块的文本
        force_semantic: 强制使用语义分块

    Returns:
        分块列表
    """
    global semantic_chunker

    # 优先使用语义分块
    if (USE_SEMANTIC_CHUNK or force_semantic) and semantic_chunker is not None:
        return semantic_chunker.split_text(text)

    # 降级到传统分块
    return split_text(text, chunk_size=300, overlap=50)


def split_text(text, chunk_size=300, overlap=50):
    """将文本切分成重叠片段（传统方式）"""
    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]

        # 尝试在句子边界处切分
        if end < len(text):
            # 查找最后一个句号、问号或感叹号
            last_period = max(
                chunk.rfind('。'),
                chunk.rfind('？'),
                chunk.rfind('！'),
                chunk.rfind('\n')
            )
            if last_period > chunk_size // 2:
                chunk = chunk[:last_period + 1]
                end = start + last_period + 1

        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap

    return chunks


def build_knowledge_base(force=False):
    """构建知识库（向量索引 + BM25索引）"""
    global chroma_client, collection

    # 多向量库模式：调用重建脚本
    if USE_MULTI_KB:
        print("\n" + "=" * 50)
        print("多向量库模式")
        print("=" * 50)
        print("\n请使用 rebuild_multi_kb.py 脚本构建多向量库:")
        print("  python rebuild_multi_kb.py")
        print("\n或使用 API 服务:")
        print("  python rag_api_server.py")
        return

    print("\n" + "=" * 50)
    print("构建知识库")
    print("=" * 50)

    # 强制重建模式：清空并重建
    if force:
        print("\n[强制重建模式]")
        # 清空BM25索引
        bm25_index.clear()

        # 尝试清空 collection 数据
        try:
            # 获取所有 ID 并删除
            results = collection.get()
            if results['ids']:
                collection.delete(ids=results['ids'])
                print(f"已清空原有数据: {len(results['ids'])} 条记录")
            else:
                print("数据库为空，开始构建")
        except Exception as e:
            # 如果出错，尝试删除并重建整个数据库目录
            print(f"清空数据失败，尝试重建数据库目录: {e}")
            import shutil
            import gc
            import time

            # 关闭连接
            try:
                del collection
                del chroma_client
            except:
                pass
            collection = None
            chroma_client = None
            gc.collect()
            time.sleep(2)

            # 删除目录
            if os.path.exists(CHROMA_DB_PATH):
                shutil.rmtree(CHROMA_DB_PATH)
                print(f"已删除旧数据库: {CHROMA_DB_PATH}")

            # 重新创建
            chroma_client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
            collection = chroma_client.get_or_create_collection(
                name="knowledge_base",
                metadata={"description": "RAG Demo 知识库"}
            )
            print("已创建新数据库")
    else:
        # 检查是否已有数据
        try:
            existing_count = collection.count()
            if existing_count > 0:
                print(f"\n知识库已有 {existing_count} 条记录")
                print("保持现有知识库 (使用 --rebuild 参数强制重建)")
                return
        except Exception as e:
            print(f"\n数据库可能已损坏，将重建: {e}")
            # 递归调用，强制重建
            return build_knowledge_base(force=True)

    # 加载文档
    print("\n加载文档...")
    documents = load_documents()

    if not documents:
        print("未找到任何文档")
        return

    # 切分并向量化
    print("\n处理文档...")
    total_chunks = 0

    # 收集所有文档用于BM25索引
    all_ids = []
    all_docs = []
    all_metas = []

    for doc in documents:
        # PDF特殊处理，支持 ODL 智能分块和传统按页切分
        if doc['type'] == 'pdf':
            doc_chunks = 0
            for page_info in doc['pages']:
                page_text = page_info['text']
                page_num = page_info['page']
                has_table = page_info.get('has_table', False)
                section = page_info.get('section', '')

                # 检查是否为 ODL 分块（智能分块，不切分）
                is_odl_chunk = page_info.get('is_odl_chunk', False)

                if is_odl_chunk:
                    # ODL 分块：直接使用，不再切分
                    chunks = [page_text]
                    section_path = page_info.get('section_path', section)
                    level = page_info.get('level', 0)
                    page_end = page_info.get('page_end', page_num)
                    bbox = page_info.get('bbox')
                else:
                    # 传统分块：按字符切分
                    chunks = split_text(page_text)
                    section_path = section
                    level = 0
                    page_end = page_num
                    bbox = None

                doc_chunks += len(chunks)

                for i, chunk in enumerate(chunks):
                    vector = embedding_model.encode(chunk).tolist()

                    # ODL 分块使用章节路径作为 ID
                    if is_odl_chunk:
                        chunk_id = f"{doc['filename']}_s{page_num}_{level}_{i}"
                    else:
                        chunk_id = f"{doc['filename']}_p{page_num}_{i}"

                    meta = {
                        'source': doc['filename'],
                        'page': page_num,
                        'page_end': page_end,
                        'chunk_index': i,
                        'has_table': has_table,
                        'section': section,
                        'section_path': section_path,
                        'level': level,
                        'is_odl_chunk': is_odl_chunk,
                        'security_level': doc.get('security_level', 'public')
                    }

                    # 添加 bounding box 信息（如果有）
                    if bbox:
                        meta['bbox'] = bbox

                    collection.add(
                        ids=[chunk_id],
                        embeddings=[vector],
                        documents=[chunk],
                        metadatas=[meta]
                    )

                    # 收集用于BM25
                    all_ids.append(chunk_id)
                    all_docs.append(chunk)
                    all_metas.append(meta)

                    total_chunks += 1
            print(f"  {doc['filename']}: {doc_chunks} 个片段")

        # Word文档处理，按内容块处理
        elif doc['type'] == 'docx':
            doc_chunks = 0
            for block in doc['blocks']:
                text = block['text']
                if len(text.strip()) < 10:  # 跳过太短的内容
                    continue

                is_table = block.get('is_table', False)
                section = block.get('section', '')

                # 表格内容整体保留，不切分
                if is_table:
                    chunks = [text]
                else:
                    chunks = split_text(text)

                for i, chunk in enumerate(chunks):
                    vector = embedding_model.encode(chunk).tolist()
                    chunk_id = f"{doc['filename']}_{doc_chunks}_{i}"

                    meta = {
                        'source': doc['filename'],
                        'chunk_index': doc_chunks,
                        'is_table': is_table,
                        'section': section,
                        'security_level': doc.get('security_level', 'public')
                    }

                    collection.add(
                        ids=[chunk_id],
                        embeddings=[vector],
                        documents=[chunk],
                        metadatas=[meta]
                    )

                    # 收集用于BM25
                    all_ids.append(chunk_id)
                    all_docs.append(chunk)
                    all_metas.append(meta)

                    total_chunks += 1
                    doc_chunks += 1
            print(f"  {doc['filename']}: {doc_chunks} 个片段")

        # Excel处理，按块处理（P1优化：智能分块）
        elif doc['type'] == 'xlsx':
            doc_chunks = 0
            for block_info in doc['rows']:
                text = block_info['text']
                if len(text.strip()) < 5:  # 跳过空内容
                    continue

                sheet = block_info['sheet']
                row_num = block_info['row']
                row_range = block_info.get('row_range', str(row_num))
                block_title = block_info.get('block_title', '')
                is_block = block_info.get('is_block', False)

                # 直接使用组合后的文本
                vector = embedding_model.encode(text).tolist()

                # 使用行范围作为ID的一部分
                chunk_id = f"{doc['filename']}_{sheet}_{row_range}"

                meta = {
                    'source': doc['filename'],
                    'sheet': sheet,
                    'row': row_num,
                    'row_range': row_range,
                    'block_title': block_title,
                    'is_block': is_block,
                    'is_excel': True,
                    'security_level': doc.get('security_level', 'public')
                }

                collection.add(
                    ids=[chunk_id],
                    embeddings=[vector],
                    documents=[text],
                    metadatas=[meta]
                )

                # 收集用于BM25
                all_ids.append(chunk_id)
                all_docs.append(text)
                all_metas.append(meta)

                doc_chunks += 1
                total_chunks += 1
            print(f"  {doc['filename']}: {doc_chunks} 个片段")

        # 其他文档类型处理（TXT等）
        else:
            chunks = split_text(doc['content'])
            print(f"  {doc['filename']}: {len(chunks)} 个片段")

            for i, chunk in enumerate(chunks):
                vector = embedding_model.encode(chunk).tolist()
                chunk_id = f"{doc['filename']}_{i}"

                meta = {
                    'source': doc['filename'],
                    'chunk_index': i,
                    'security_level': doc.get('security_level', 'public')
                }

                collection.add(
                    ids=[chunk_id],
                    embeddings=[vector],
                    documents=[chunk],
                    metadatas=[meta]
                )

                # 收集用于BM25
                all_ids.append(chunk_id)
                all_docs.append(chunk)
                all_metas.append(meta)

                total_chunks += 1

    # 构建BM25索引 (P4优化)
    if USE_HYBRID_SEARCH and all_docs:
        print("\n构建BM25索引...")
        bm25_index.add_documents(all_ids, all_docs, all_metas)
        bm25_index.save(BM25_INDEX_PATH)
        print(f"      BM25索引构建完成: {len(all_docs)} 个文档")

    print(f"\n知识库构建完成，共 {total_chunks} 个片段")
    print("\n提示: 图谱构建请运行 'python graph_build.py --help'")


# ========== 检索函数 ==========
def reciprocal_rank_fusion(results_list, weights=None, k=60):
    """
    RRF (Reciprocal Rank Fusion) 算法 - 融合多个检索结果

    参数:
        results_list: 多个检索结果的列表 [{'ids': [], 'documents': [], 'metadatas': [], 'distances': []}, ...]
        weights: 每个检索器的权重，默认等权重
        k: RRF参数，防止除零并平滑排名

    返回:
        融合后的结果
    """
    if not results_list:
        return {'ids': [[]], 'documents': [[]], 'metadatas': [[]], 'distances': [[]]}

    if weights is None:
        weights = [1.0] * len(results_list)

    # 文档分数累加器
    doc_scores = {}  # {doc_id: {'score': float, 'doc': str, 'meta': dict}}

    for results, weight in zip(results_list, weights):
        if not results['documents'][0]:
            continue

        for rank, (doc_id, doc, meta) in enumerate(zip(
            results['ids'][0],
            results['documents'][0],
            results['metadatas'][0]
        )):
            # RRF公式: score += weight / (k + rank + 1)
            rrf_score = weight / (k + rank + 1)

            if doc_id not in doc_scores:
                doc_scores[doc_id] = {
                    'score': 0.0,
                    'doc': doc,
                    'meta': meta
                }
            doc_scores[doc_id]['score'] += rrf_score

    # 按分数排序
    sorted_items = sorted(doc_scores.items(), key=lambda x: x[1]['score'], reverse=True)

    return {
        'ids': [[item[0] for item in sorted_items]],
        'documents': [[item[1]['doc'] for item in sorted_items]],
        'metadatas': [[item[1]['meta'] for item in sorted_items]],
        'distances': [[item[1]['score'] for item in sorted_items]]
    }


def rerank_results(query, results, top_k=5):
    """
    对检索结果进行重排序 (P3优化)

    使用CrossEncoder模型计算query与每个文档的精确相关性分数，
    重新排序后返回top_k个最相关的结果。
    """
    if not reranker or not results['documents'][0]:
        return results

    # 构建query-doc对
    pairs = [(query, doc) for doc in results['documents'][0]]

    # 计算重排序分数
    scores = reranker.predict(pairs)

    # 按分数排序
    sorted_indices = np.argsort(scores)[::-1]  # 降序

    # 重新组织结果
    reranked_results = {
        'ids': [[results['ids'][0][i] for i in sorted_indices[:top_k]]],
        'documents': [[results['documents'][0][i] for i in sorted_indices[:top_k]]],
        'metadatas': [[results['metadatas'][0][i] for i in sorted_indices[:top_k]]],
        'distances': [[float(scores[i]) for i in sorted_indices[:top_k]]]
    }

    return reranked_results


def search_knowledge(query, top_k=5, allowed_levels=None, role=None, department=None, collections=None):
    """
    混合检索 (P4优化: 向量检索 + BM25 + RRF融合 + Rerank)

    检索流程:
    1. 向量检索 - 语义相似度
    2. BM25检索 - 关键词匹配
    3. RRF融合 - 合并两种检索结果
    4. Rerank精排 - 最终排序

    Args:
        query: 查询文本
        top_k: 返回结果数
        allowed_levels: 允许访问的安全级别列表（单向量库模式）
        role: 用户角色（多向量库模式）
        department: 用户部门（多向量库模式）
        collections: 指定向量库列表（多向量库模式，可选）
    """
    results_list = []
    weights = []

    # 多向量库模式
    if USE_MULTI_KB and kb_manager:
        return _search_multi_kb(query, top_k, role, department, collections)

    # 单向量库模式（原有逻辑）
    where_filter = None
    if allowed_levels:
        where_filter = {"security_level": {"$in": allowed_levels}}

    # 1. 向量检索
    query_vector = embedding_model.encode(query).tolist()
    recall_k = RERANK_CANDIDATES if (USE_RERANK or USE_HYBRID_SEARCH) else top_k

    query_kwargs = {
        "query_embeddings": [query_vector],
        "n_results": recall_k
    }
    if where_filter:
        query_kwargs["where"] = where_filter

    vector_results = collection.query(**query_kwargs)
    results_list.append(vector_results)
    weights.append(VECTOR_WEIGHT)

    # 2. BM25检索 (如果启用混合检索)
    if USE_HYBRID_SEARCH and bm25_index.bm25:
        bm25_results = bm25_index.search(query, top_k=recall_k)
        # BM25 后处理：过滤无权限文档
        if where_filter and bm25_results['metadatas'][0]:
            allowed_set = set(allowed_levels)
            filtered_ids = []
            filtered_docs = []
            filtered_metas = []
            filtered_scores = []
            for bid, bdoc, bmeta, bscore in zip(
                bm25_results['ids'][0], bm25_results['documents'][0],
                bm25_results['metadatas'][0], bm25_results['distances'][0]
            ):
                doc_level = bmeta.get('security_level', 'public')
                if doc_level in allowed_set:
                    filtered_ids.append(bid)
                    filtered_docs.append(bdoc)
                    filtered_metas.append(bmeta)
                    filtered_scores.append(bscore)
            bm25_results = {
                'ids': [filtered_ids],
                'documents': [filtered_docs],
                'metadatas': [filtered_metas],
                'distances': [filtered_scores]
            }
        results_list.append(bm25_results)
        weights.append(BM25_WEIGHT)

    # 3. RRF融合
    if len(results_list) > 1:
        fused_results = reciprocal_rank_fusion(results_list, weights)
    else:
        fused_results = results_list[0]

    # 4. Rerank精排
    if USE_RERANK and reranker:
        fused_results = rerank_results(query, fused_results, top_k)
    else:
        # 截取top_k
        fused_results = {
            'ids': [fused_results['ids'][0][:top_k]],
            'documents': [fused_results['documents'][0][:top_k]],
            'metadatas': [fused_results['metadatas'][0][:top_k]],
            'distances': [fused_results['distances'][0][:top_k]]
        }

    return fused_results


def _search_multi_kb(query, top_k=5, role=None, department=None, target_collections=None):
    """
    多向量库检索

    Args:
        query: 查询文本
        top_k: 返回结果数
        role: 用户角色
        department: 用户部门
        target_collections: 指定向量库列表（可选，不传则自动路由）
    """
    # 确定目标向量库
    if target_collections is None:
        if role and department:
            # 使用路由器自动选择
            accessible = get_accessible_collections(role, department, 'read')
            target_collections = kb_router.route(query, role, department, accessible)
        else:
            # 默认只查 public_kb
            target_collections = ['public_kb']

    if not target_collections:
        return {'ids': [[]], 'documents': [[]], 'metadatas': [[]], 'distances': [[]]}

    # 生成查询向量
    query_vector = embedding_model.encode(query).tolist()
    recall_k = RERANK_CANDIDATES if (USE_RERANK or USE_HYBRID_SEARCH) else top_k

    # 并行检索多个向量库
    all_results = []
    for coll_name in target_collections:
        try:
            coll = kb_manager.get_collection(coll_name)
            if coll is None:
                continue

            # 向量检索
            results = coll.query(
                query_embeddings=[query_vector],
                n_results=recall_k
            )

            # 添加来源标记
            if results['metadatas'] and results['metadatas'][0]:
                for meta in results['metadatas'][0]:
                    meta['_collection'] = coll_name

            all_results.append(results)

            # BM25 检索（如果启用）
            if USE_HYBRID_SEARCH:
                try:
                    bm25 = kb_manager.get_bm25_index(coll_name)
                    if bm25.bm25:
                        bm25_results = bm25.search(query, top_k=recall_k)
                        if bm25_results['metadatas'] and bm25_results['metadatas'][0]:
                            for meta in bm25_results['metadatas'][0]:
                                meta['_collection'] = coll_name
                        all_results.append(bm25_results)
                except Exception as e:
                    pass  # BM25 索引可能不存在

        except Exception as e:
            print(f"检索向量库 {coll_name} 失败: {e}")
            continue

    if not all_results:
        return {'ids': [[]], 'documents': [[]], 'metadatas': [[]], 'distances': [[]]}

    # 合并结果
    if len(all_results) == 1:
        fused_results = all_results[0]
    else:
        # 使用 RRF 融合
        weights = [VECTOR_WEIGHT if i % 2 == 0 else BM25_WEIGHT for i in range(len(all_results))]
        fused_results = reciprocal_rank_fusion(all_results, weights)

    # Rerank 精排
    if USE_RERANK and reranker:
        fused_results = rerank_results(query, fused_results, top_k)
    else:
        fused_results = {
            'ids': [fused_results['ids'][0][:top_k]],
            'documents': [fused_results['documents'][0][:top_k]],
            'metadatas': [fused_results['metadatas'][0][:top_k]],
            'distances': [fused_results['distances'][0][:top_k]]
        }

    return fused_results


def check_restricted_documents(query, allowed_levels, top_k=3, role=None, department=None):
    """
    检测是否存在超出用户权限的相关文档

    当用户查询没有返回结果时，用于判断是因为知识库中确实没有相关信息，
    还是因为相关信息存在于用户无权访问的文档中。

    Args:
        query: 查询文本
        allowed_levels: 用户当前允许访问的安全级别列表（单向量库模式）
        top_k: 检索数量
        role: 用户角色（多向量库模式）
        department: 用户部门（多向量库模式）

    Returns:
        {
            "has_restricted": bool,
            "restricted_levels": list,
            "restricted_sources": list,
            "top_restricted_score": float
        }
    """
    # 多向量库模式
    if USE_MULTI_KB and kb_manager and role and department:
        return _check_restricted_multi_kb(query, role, department, top_k)

    # 单向量库模式
    if not allowed_levels:
        return {"has_restricted": False, "restricted_levels": [], "restricted_sources": [], "top_restricted_score": 0.0}

    all_levels = {"public", "internal", "confidential", "secret"}
    allowed_set = set(allowed_levels)
    restricted_levels = all_levels - allowed_set

    if not restricted_levels:
        return {"has_restricted": False, "restricted_levels": [], "restricted_sources": [], "top_restricted_score": 0.0}

    restricted_filter = {"security_level": {"$in": list(restricted_levels)}}

    query_vector = embedding_model.encode(query).tolist()
    query_kwargs = {
        "query_embeddings": [query_vector],
        "n_results": top_k,
        "where": restricted_filter
    }

    try:
        restricted_results = collection.query(**query_kwargs)

        docs = restricted_results.get('documents', [[]])[0]
        metas = restricted_results.get('metadatas', [[]])[0]
        distances = restricted_results.get('distances', [[]])[0]

        if not docs:
            return {"has_restricted": False, "restricted_levels": [], "restricted_sources": [], "top_restricted_score": 0.0}

        found_levels = set()
        found_sources = set()
        top_score = 0.0

        for doc, meta, dist in zip(docs, metas, distances):
            level = meta.get('security_level', 'public')
            source = meta.get('source', '未知')
            found_levels.add(level)
            found_sources.add(source)
            if dist > top_score:
                top_score = dist

        level_names = {
            "public": "公开",
            "internal": "内部",
            "confidential": "机密",
            "secret": "绝密"
        }

        return {
            "has_restricted": True,
            "restricted_levels": [level_names.get(l, l) for l in found_levels],
            "restricted_sources": list(found_sources)[:3],
            "top_restricted_score": top_score
        }

    except Exception as e:
        print(f"检测受限文档时出错: {e}")
        return {"has_restricted": False, "restricted_levels": [], "restricted_sources": [], "top_restricted_score": 0.0}


def _check_restricted_multi_kb(query, role, department, top_k=3):
    """多向量库模式下的受限文档检测"""
    from auth.gateway import get_accessible_collections

    # 获取所有向量库和可访问向量库
    all_collections = [c.name for c in kb_manager.list_collections()]
    accessible = set(get_accessible_collections(role, department, 'read'))
    restricted = set(all_collections) - accessible

    if not restricted:
        return {"has_restricted": False, "restricted_levels": [], "restricted_sources": [], "top_restricted_score": 0.0}

    query_vector = embedding_model.encode(query).tolist()
    found_sources = set()
    top_score = 0.0

    for coll_name in restricted:
        try:
            coll = kb_manager.get_collection(coll_name)
            if coll is None:
                continue

            results = coll.query(
                query_embeddings=[query_vector],
                n_results=top_k
            )

            if results['metadatas'] and results['metadatas'][0]:
                for meta in results['metadatas'][0]:
                    source = meta.get('source', '未知')
                    found_sources.add(source)
                if results['distances'] and results['distances'][0]:
                    for dist in results['distances'][0]:
                        if dist > top_score:
                            top_score = dist
        except Exception:
            continue

    return {
        "has_restricted": len(found_sources) > 0,
        "restricted_levels": [coll.replace('dept_', '') for coll in restricted if coll in found_sources or True][:3],
        "restricted_sources": list(found_sources)[:3],
        "top_restricted_score": top_score
    }


def aggregate_excel_rows(results, max_rows=10):
    """
    聚合Excel相邻行数据，形成更完整的表格上下文

    当检索结果包含Excel数据时，将同一工作表的相邻行合并展示，
    让大模型能看到完整的表格结构，而非碎片化的单行数据。
    """
    if not results['documents'][0]:
        return results

    # 按文件和工作表分组
    excel_groups = {}  # {(source, sheet): [rows]}
    other_docs = []    # 非Excel文档

    for i, (doc, meta) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
        source = meta.get('source', '')

        # 检查是否为Excel文件
        if source.endswith('.xlsx') and 'sheet' in meta and 'row' in meta:
            key = (source, meta['sheet'])
            if key not in excel_groups:
                excel_groups[key] = []
            excel_groups[key].append({
                'row': meta['row'],
                'doc': doc,
                'meta': meta,
                'original_index': i
            })
        else:
            other_docs.append({
                'doc': doc,
                'meta': meta,
                'original_index': i
            })

    # 如果没有Excel数据，直接返回原结果
    if not excel_groups:
        return results

    # 聚合Excel数据
    aggregated_excel = []
    for (source, sheet), rows in excel_groups.items():
        # 按行号排序
        rows.sort(key=lambda x: x['row'])

        # 获取表头信息
        header = None
        for r in rows:
            if r['meta'].get('is_header'):
                header = r['doc']
                break

        # 构建聚合后的内容
        if len(rows) == 1:
            # 只有一行，直接使用
            aggregated_doc = rows[0]['doc']
        else:
            # 多行，聚合展示
            rows_text = []
            for r in rows[:max_rows]:
                # 去掉可能重复的表头信息
                text = r['doc']
                if text.startswith('【表头:'):
                    # 提取实际内容
                    parts = text.split('\n', 1)
                    if len(parts) > 1:
                        text = parts[1]
                rows_text.append(f"第{r['row']}行: {text}")

            aggregated_doc = f"【Excel表格数据 - {sheet}】\n" + "\n".join(rows_text)
            if len(rows) > max_rows:
                aggregated_doc += f"\n... (共{len(rows)}行，已显示前{max_rows}行)"

        aggregated_excel.append({
            'doc': aggregated_doc,
            'meta': {
                'source': source,
                'sheet': sheet,
                'row_range': f"{rows[0]['row']}-{rows[-1]['row']}",
                'is_excel': True
            },
            'original_index': rows[0]['original_index']
        })

    # 合并Excel和非Excel结果，保持原顺序
    all_results = aggregated_excel + other_docs
    all_results.sort(key=lambda x: x['original_index'])

    # 重建results结构
    new_documents = [r['doc'] for r in all_results]
    new_metadatas = [r['meta'] for r in all_results]
    new_distances = [results['distances'][0][r['original_index']] for r in all_results] if results.get('distances') else [0] * len(all_results)

    return {
        'ids': [results['ids'][0][r['original_index']] for r in all_results],
        'documents': [new_documents],
        'metadatas': [new_metadatas],
        'distances': [new_distances]
    }


def generate_answer(query, context):
    """调用大模型生成回答（P2优化：添加置信度标注）"""
    prompt = f"""你是一个严谨的智能助手，请根据以下参考资料回答用户的问题。

【严格约束】
1. 只能基于【参考资料】中的信息回答，禁止使用你的先验知识
2. 若参考资料中没有答案，直接回复"参考资料中未找到相关信息"
3. 不要推测、不要补充、不要编造
4. 必须标注信息来源（文件名、页码/行号等）

【回答格式】
1. 直接回答问题，用结构化方式呈现（表格、列表等）
2. 标注信息来源
3. 在回答末尾添加【置信度评估】

【置信度评估标准】
- 高：多个来源一致，信息完整，直接命中关键词
- 中：信息部分匹配，需要一定推理，来源较少
- 低：信息模糊，需要较多推理，来源单一或存在矛盾

【置信度格式】
---
置信度：高/中/低
评估理由：简要说明为什么给出这个置信度

参考资料：
{context}

用户问题：{query}

请回答："""

    try:
        response = llm_client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1500
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"调用大模型失败: {str(e)}"


def chat(query=None):
    """交互式问答或单次问答"""
    # 多向量库模式下的用户身份
    current_role = 'admin'
    current_department = ''

    if query:
        # 单次问答模式
        return process_query(query, role=current_role, department=current_department)

    # 交互模式
    print("\n" + "=" * 50)
    print("知识库问答")
    if USE_MULTI_KB:
        print("(多向量库模式)")
    print("=" * 50)
    print("命令 (以 / 开头):")
    print("  /quit    - 退出程序")
    print("  /rebuild - 完全重建知识库")
    print("  /sync    - 同步文档（增量更新）")
    print("  /add <文件名> - 添加单个文件")
    print("  /del <文件名> - 删除单个文件")
    print("  /list    - 列出已索引的文件")
    print("  /help    - 显示帮助信息")
    if USE_MULTI_KB:
        print("  /role <role> - 设置角色 (admin/manager/user)")
        print("  /dept <dept> - 设置部门 (finance/hr/tech/...)")
    print("其他输入将作为问题进行问答")
    print("=" * 50)
    if USE_MULTI_KB:
        print(f"当前身份: {current_role}/{current_department or '全部'}")

    while True:
        print("\n" + "-" * 30)
        query = input("\n请输入问题或命令: ").strip()

        if not query:
            continue

        # 解析命令（以 / 开头）
        if query.startswith('/'):
            cmd = query.lower()

            if cmd == '/quit':
                print("\n感谢使用，再见!")
                break

            if cmd == '/rebuild':
                build_knowledge_base(force=True)
                continue

            if cmd == '/sync':
                sync_documents()
                continue

            if cmd == '/list':
                if USE_MULTI_KB:
                    print(f"\n向量库列表:")
                    for coll_info in kb_manager.list_collections():
                        print(f"  {coll_info.name}: {coll_info.document_count} 条记录")
                else:
                    files = list_indexed_files()
                    print(f"\n已索引文件 ({len(files)} 个):")
                    for f, count in sorted(files.items()):
                        print(f"  {f}: {count} 片段")
                continue

            if cmd == '/help':
                print("\n命令列表:")
                print("  /quit    - 退出程序")
                print("  /rebuild - 完全重建知识库")
                print("  /sync    - 同步文档（增量更新）")
                print("  /add <文件名> - 添加单个文件")
                print("  /del <文件名> - 删除单个文件")
                print("  /list    - 列出已索引的文件")
                print("  /help    - 显示帮助信息")
                if USE_MULTI_KB:
                    print("  /role <role> - 设置角色 (admin/manager/user)")
                    print("  /dept <dept> - 设置部门 (finance/hr/tech/...)")
                continue

            if cmd.startswith('/add '):
                filename = query[5:].strip()
                filepath = os.path.join(DOCUMENTS_PATH, filename)
                if os.path.exists(filepath):
                    add_file_to_index(filepath)
                else:
                    print(f"文件不存在: {filename}")
                continue

            if cmd.startswith('/del '):
                filename = query[5:].strip()
                delete_file_from_index(filename)
                continue

            # 多向量库模式下的角色和部门设置
            if USE_MULTI_KB and cmd.startswith('/role '):
                new_role = query[6:].strip().lower()
                if new_role in ['admin', 'manager', 'user']:
                    current_role = new_role
                    print(f"角色已设置为: {current_role}")
                    print(f"当前身份: {current_role}/{current_department or '全部'}")
                else:
                    print(f"无效角色: {new_role}，有效值: admin, manager, user")
                continue

            if USE_MULTI_KB and cmd.startswith('/dept '):
                new_dept = query[6:].strip().lower()
                current_department = new_dept
                print(f"部门已设置为: {current_department}")
                print(f"当前身份: {current_role}/{current_department or '全部'}")
                continue

            # 未知命令
            print(f"未知命令: {query}")
            print("输入 /help 查看可用命令")
            continue

        # 不是命令，作为问题处理
        process_query(query, role=current_role, department=current_department)


def process_query(query, role=None, department=None):
    """处理单个查询

    Args:
        query: 查询文本
        role: 用户角色（多向量库模式）
        department: 用户部门（多向量库模式）
    """
    # 检索相关内容
    print("\n[检索中...]")

    # 多向量库模式
    if USE_MULTI_KB and role and department:
        results = search_knowledge(query, top_k=5, role=role, department=department)
    else:
        results = search_knowledge(query, top_k=5)

    if not results['documents'][0]:
        print("未找到相关内容")
        return "未找到相关内容"

    # 聚合Excel数据（P1优化）
    results = aggregate_excel_rows(results)

    # 组装上下文
    context_parts = []
    sources = set()

    for doc, meta in zip(results['documents'][0], results['metadatas'][0]):
        # 构建来源信息字符串
        source_parts = [meta['source']]

        if 'page' in meta:
            source_parts.append(f"第{meta['page']}页")
        if 'sheet' in meta and 'row' in meta:
            source_parts.append(f"工作表\"{meta['sheet']}\"第{meta['row']}行")
        if 'row_range' in meta:
            source_parts.append(f"第{meta['row_range']}行")
        if 'section' in meta and meta['section']:
            source_parts.append(f"【{meta['section']}】")
        if meta.get('is_table'):
            source_parts.append("(表格)")
        if meta.get('is_header'):
            source_parts.append("(表头)")
        if meta.get('is_excel'):
            source_parts.append("(Excel数据)")

        source_str = " ".join(source_parts)
        context_parts.append(f"【来源: {source_str}】\n{doc}")
        sources.add(meta['source'])

    context = "\n\n".join(context_parts)

    print(f"找到 {len(results['documents'][0])} 个相关片段")
    print(f"涉及文档: {', '.join(sources)}")

    # 生成回答
    print("\n[生成回答中...]")
    answer = generate_answer(query, context)

    print("\n" + "=" * 50)
    print("回答:")
    print("-" * 30)
    print(answer)
    print("-" * 30)

    # 显示检索到的片段
    print("\n参考片段:")
    for i, (doc, meta) in enumerate(zip(results['documents'][0], results['metadatas'][0]), 1):
        preview = doc[:100] + "..." if len(doc) > 100 else doc

        # 构建来源信息
        source_info = meta['source']
        if 'page' in meta:
            source_info += f" 第{meta['page']}页"
        if 'sheet' in meta and 'row' in meta:
            source_info += f" [{meta['sheet']} 第{meta['row']}行]"
        if 'row_range' in meta:
            source_info += f" [{meta['sheet']} 第{meta['row_range']}行]"
        if 'section' in meta and meta['section']:
            source_info += f" 【{meta['section']}】"
        if meta.get('is_table'):
            source_info += " [表格]"
        if meta.get('is_header'):
            source_info += " [表头]"
        if meta.get('is_excel'):
            source_info += " [Excel]"

        print(f"  [{i}] {source_info}: {preview}")

    return answer


# ========== 主程序 ==========
if __name__ == "__main__":
    import sys

    # 解析命令行参数
    args = sys.argv[1:]
    force_rebuild = "--rebuild" in args
    sync_mode = "--sync" in args
    list_mode = "--list" in args

    if force_rebuild:
        args.remove("--rebuild")
    if sync_mode:
        args.remove("--sync")
    if list_mode:
        args.remove("--list")

    # 命令行模式
    if list_mode:
        # 列出文件
        files = list_indexed_files()
        print(f"已索引文件 ({len(files)} 个):")
        for f, count in sorted(files.items()):
            print(f"  {f}: {count} 片段")

    elif sync_mode:
        # 同步文档
        sync_documents()

    elif force_rebuild:
        # 强制重建
        build_knowledge_base(force=True)
        print(f"\n知识库统计:")
        print(f"  总片段数: {collection.count()}")

    elif args:
        # 单次问答模式
        build_knowledge_base(force=False)
        query = " ".join(args)
        print(f"\n问题: {query}")
        process_query(query)

    else:
        # 交互模式
        print(f"\n当前知识库: {collection.count()} 个片段")
        chat()

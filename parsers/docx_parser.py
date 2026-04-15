"""
Word 传统解析器

使用 docx2txt + python-docx 提取 Word 文档内容。
作为 Docling 解析器的回退方案。
"""

import docx2txt
from docx import Document


def extract_text_from_docx_traditional(filepath):
    """
    Word 传统解析（简化版，作为 Docling 的回退）

    保留基础文本提取和简单表格处理
    """
    content_blocks = []

    # 首先尝试用 docx2txt 提取（兼容性更好）
    try:
        full_text = docx2txt.process(filepath)
        if full_text and full_text.strip():
            for para in full_text.split('\n\n'):
                text = para.strip()
                if text:
                    content_blocks.append({
                        'text': text,
                        'is_heading': False,
                        'section': text[:30] if len(text) < 30 else text[:30] + '...',
                        'is_table': False
                    })
            return content_blocks
    except Exception as e:
        print(f"      docx2txt 解析失败，尝试 python-docx: {e}")

    # 备用：使用 python-docx
    try:
        doc = Document(filepath)
        current_section = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 简单标题检测
            is_heading = para.style.name.startswith(('Heading', '标题'))
            if is_heading:
                current_section = text

            content_blocks.append({
                'text': text,
                'is_heading': is_heading,
                'section': current_section if current_section else text[:20],
                'is_table': False
            })

        # 简单表格提取
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    table_rows.append(row_text)
            if table_rows:
                content_blocks.append({
                    'text': "【表格】\n" + "\n".join(table_rows),
                    'is_heading': False,
                    'section': current_section,
                    'is_table': True
                })

    except Exception as e:
        print(f"      Word 解析错误 {filepath}: {e}")
    return content_blocks
